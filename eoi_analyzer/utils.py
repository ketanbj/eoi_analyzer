from pathlib import Path
from typing import Iterable, List

from .config import SUPPORTED_DOCUMENT_EXTENSIONS
from .exceptions import DocumentExtractionError, EmptyDocumentError, UnsupportedDocumentError

def extract_text_from_pdf(pdf_path):
    """
    Extract text from a PDF document.
    """
    import pdfplumber

    extracted_text = ""
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                extracted_text += text + "\n"
    return extracted_text


def extract_text_from_text_file(path: Path) -> str:
    """
    Extract text from plain text and markdown files.
    """
    return path.read_text(encoding="utf-8")


def extract_text_from_docx(path: Path) -> str:
    """
    Extract paragraph and table text from a DOCX document.
    """
    try:
        from docx import Document
    except ImportError as exc:
        raise DocumentExtractionError(
            "DOCX support requires the python-docx dependency."
        ) from exc

    document = Document(path)
    parts: List[str] = [paragraph.text for paragraph in document.paragraphs if paragraph.text]

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def normalize_extracted_text(text: str, source_path: Path) -> str:
    """
    Normalize extracted document text and fail early when no usable text is found.
    """
    normalized = "\n".join(line.rstrip() for line in text.splitlines()).strip()
    if not normalized:
        raise EmptyDocumentError(f"No extractable text found in {source_path}.")
    return normalized


def extract_text_from_document(document_path) -> str:
    """
    Extract text from a supported document type.
    """
    path = Path(document_path)
    suffix = path.suffix.lower()

    try:
        if suffix == ".pdf":
            text = extract_text_from_pdf(path)
        elif suffix in {".txt", ".md", ".markdown"}:
            text = extract_text_from_text_file(path)
        elif suffix == ".docx":
            text = extract_text_from_docx(path)
        else:
            raise UnsupportedDocumentError(f"Unsupported document type: {path.suffix}")
    except (UnsupportedDocumentError, EmptyDocumentError):
        raise
    except Exception as exc:
        raise DocumentExtractionError(f"Could not extract text from {path}: {exc}") from exc

    return normalize_extracted_text(text, path)


def is_supported_document(path: Path, extensions=None) -> bool:
    """
    Return True when a path points to a supported document type.
    """
    allowed_extensions = extensions or SUPPORTED_DOCUMENT_EXTENSIONS
    return path.is_file() and path.suffix.lower() in allowed_extensions


def iter_document_paths(folder_path, recursive: bool = True, extensions=None) -> Iterable[Path]:
    """
    Yield supported document files in a folder in deterministic order.
    """
    folder = Path(folder_path)
    if not folder.exists():
        raise FileNotFoundError(f"Folder does not exist: {folder}")
    if not folder.is_dir():
        raise NotADirectoryError(f"Expected a folder path, got: {folder}")

    pattern = "**/*" if recursive else "*"
    paths = sorted(path for path in folder.glob(pattern) if is_supported_document(path, extensions=extensions))
    return paths
